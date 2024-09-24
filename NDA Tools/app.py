import os
from flask import Flask, render_template, request, redirect, url_for, send_file, flash
from werkzeug.utils import secure_filename
import openai
import docx
from io import BytesIO
from docx.enum.text import WD_COLOR_INDEX  # Importando o WD_COLOR_INDEX corretamente
from docx.enum.text import WD_ALIGN_PARAGRAPH  # Alinhamento justificado do docx
from docx.shared import Inches  # Para definir o tamanho da imagem
from docx.shared import Pt
import pytesseract
from PIL import Image
from pdf2image import convert_from_path
from docx.shared import Cm
import zipfile

app = Flask(__name__)

# Configurações do diretório de upload
UPLOAD_FOLDER = 'uploads'
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.secret_key = 'supersecretkey'

# Configure sua chave de API do OpenAI
openai.api_key = 'SUA CHAVE API AQUI'

# Configuração do caminho do Tesseract (modifique conforme necessário)
pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'

# Função para ler o conteúdo do arquivo DOCX
def read_docx(file_path):
    doc = docx.Document(file_path)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    return '\n'.join(full_text)

# Função para realizar OCR em uma imagem
def ocr_image(file_path):
    img = Image.open(file_path)
    text = pytesseract.image_to_string(img)
    return text

# Função para realizar OCR em PDFs
def ocr_pdf(file_path):
    images = convert_from_path(file_path)
    ocr_text = ""
    for img in images:
        ocr_text += pytesseract.image_to_string(img)
    return ocr_text

# Função para realçar o texto no DOCX
def highlight_text(paragraph, text_to_highlight):
    for run in paragraph.runs:
        if text_to_highlight in run.text:
            run.font.highlight_color = WD_COLOR_INDEX.YELLOW  # Realça com a cor amarela

# Função para aplicar a fonte e o tamanho
def set_font(run, font_name='Arial Nova Light', size=10):
    run.font.name = font_name
    run.font.size = Pt(size)

# Função para extrair imagens do arquivo DOCX (extração direta do zip)
def extract_images_from_docx(file_path):
    images = []
    with zipfile.ZipFile(file_path, 'r') as docx_file:
        for file in docx_file.namelist():
            if file.startswith('word/media/'):
                # Extraímos e salvamos cada imagem na pasta temporária
                image_data = docx_file.read(file)
                image_filename = os.path.join(app.config['UPLOAD_FOLDER'], file.split('/')[-1])
                with open(image_filename, 'wb') as img_file:
                    img_file.write(image_data)
                    images.append(image_filename)
    return images

# Rota principal para carregar a página HTML
@app.route('/')
def index():
    return render_template('index.html')

# Rota para processar os arquivos e realizar a comparação
@app.route('/compare', methods=['POST'])
def compare_ndas():
    if 'standard_nda' not in request.files or 'user_nda' not in request.files:
        flash('Ambos os arquivos são obrigatórios para a comparação!', 'error')
        return redirect(url_for('index'))

    standard_nda_file = request.files['standard_nda']
    user_nda_file = request.files['user_nda']

    if standard_nda_file.filename == '' or user_nda_file.filename == '':
        flash('Por favor, selecione os dois arquivos para comparar.', 'error')
        return redirect(url_for('index'))

    # Salvar os arquivos
    standard_nda_path = os.path.join(app.config['UPLOAD_FOLDER'], secure_filename(standard_nda_file.filename))
    user_nda_path = os.path.join(app.config['UPLOAD_FOLDER'], secure_filename(user_nda_file.filename))
    standard_nda_file.save(standard_nda_path)
    user_nda_file.save(user_nda_path)

    # Ler o conteúdo dos arquivos (usa OCR para PDFs e leitura normal para DOCX)
    if standard_nda_file.filename.endswith('.pdf'):
        standard_nda_content = ocr_pdf(standard_nda_path)
    else:
        standard_nda_content = read_docx(standard_nda_path)
        standard_nda_doc = docx.Document(standard_nda_path)

    if user_nda_file.filename.endswith('.pdf'):
        user_nda_content = ocr_pdf(user_nda_path)
    else:
        user_nda_content = read_docx(user_nda_path)

    # Extração de imagens do documento "standard_nda"
    images = extract_images_from_docx(standard_nda_path)

    # Chamada para a API da OpenAI para comparar os NDAs e realçar as modificações ///// AJUSTE CONFORME A SUA NECESSIDADE /////
    messages = [
        {"role": "system", "content": "Você é um especialista em análise de NDAs."},
        #{"role": "user", "content": f"Compare o NDA padrão com o NDA do cliente para gerar um novo NDA ajustado, utilize estes dados para preencher todos os campos vazios, Não deixe campos com XXXXXXXX ou ____________, utilize cláusula multas do NDA Cliente,  Destaque as palavras modificadas no novo documento em amarelo e, em seguida, apresente primeiro o NDA ajustado e depois uma lista clara das inconsistências encontradas.\n\nNDA Padrão:\n{standard_nda_content}\n\nNDA Cliente:\n{user_nda_content}"}
        #{"role": "user", "content": f"Compare o NDA padrão com o NDA do cliente para gerar um novo NDA ajustado. Preencha todos os campos vazios e **não deixe campos com XXXXXXXX ou ____________.** **se atente aqui substituindo multas do NDA padrão por multas do NDA do cliente.** Destaque todas as palavras e seções modificadas no novo documento em amarelo. Após apresentar o NDA ajustado, forneça uma lista clara de todas as inconsistências encontradas entre os dois documentos.\n\nNDA Padrão:\n{standard_nda_content}\n\nNDA Cliente:\n{user_nda_content}"}
        # SEM VALOR DE MULTAS {"role": "user", "content": f"Compare o NDA padrão com o NDA do cliente para gerar um novo NDA ajustado. Preencha todos os campos vazios e **não deixe campos com XXXXXXXX ou ____________.** **se atente aqui remova multas do NDA padrão.** Destaque todas as palavras e seções modificadas no novo documento em amarelo. Após apresentar o NDA ajustado, forneça uma lista clara de todas as inconsistências encontradas entre os dois documentos.\n\nNDA Padrão:\n{standard_nda_content}\n\nNDA Cliente:\n{user_nda_content}"}
        # REMOVER MULTAS {"role": "user", "content": f"Compare o NDA padrão com o NDA do cliente para gerar um novo NDA ajustado. Preencha todos os campos vazios e **não deixe campos com XXXXXXXX ou ____________.** **se atente aqui remova multas do NDA padrão.** Destaque todas as palavras e seções modificadas no novo documento em amarelo. Após apresentar o NDA ajustado, forneça uma lista clara de todas as inconsistências encontradas entre os dois documentos.\n\nNDA Padrão:\n{standard_nda_content}\n\nNDA Cliente:\n{user_nda_content}"}
        # IGNORA MULTAS {"role": "user", "content": f"Compare o NDA padrão com o NDA do cliente para gerar um novo NDA ajustado. Preencha todos os campos vazios e **não deixe campos com XXXXXXXX ou ____________.** **se atente aqui ignore multas do NDA padrão.** Destaque todas as palavras e seções modificadas no novo documento em amarelo. Após apresentar o NDA ajustado, forneça uma lista clara de todas as inconsistências encontradas entre os dois documentos.\n\nNDA Padrão:\n{standard_nda_content}\n\nNDA Cliente:\n{user_nda_content}"}
        # SUBSTITUI CLÁUSULAS {"role": "user", "content": f"Compare o NDA padrão com o NDA do cliente para gerar um novo NDA ajustado. Preencha todos os campos vazios e **não deixe campos com XXXXXXXX ou ____________.** **se atente aqui substitua todas inconsistências do NDA padrão por cláusulas do NDA cliente.** Destaque todas as palavras e seções modificadas no novo documento em amarelo. Após apresentar o NDA ajustado, forneça uma lista clara de todas as inconsistências encontradas entre os dois documentos.\n\nNDA Padrão:\n{standard_nda_content}\n\nNDA Cliente:\n{user_nda_content}"}
        # SUBSTITUI INCONSISTÊNCIAS {"role": "user", "content": f"Compare o NDA padrão com o NDA do cliente para gerar um novo NDA ajustado. Preencha todos os campos vazios e **não deixe campos com XXXXXXXX ou ____________.** **Ao encontrar inconsistências no NDA do cliente, substitua as cláusulas inconsistentes pelas cláusulas equivalentes do NDA padrão.** Destaque todas as palavras e seções modificadas no novo documento em amarelo. Após apresentar o NDA ajustado, forneça uma lista clara de todas as inconsistências encontradas entre os dois documentos.\n\nNDA Padrão:\n{standard_nda_content}\n\nNDA Cliente:\n{user_nda_content}"}
        # SUBSTITUI INCONSISTÊNCIAS E REMOVE MULTAS {"role": "user", "content": f"Compare o NDA padrão com o NDA do cliente para gerar um novo NDA ajustado. Preencha todos os campos vazios e **não deixe campos com XXXXXXXX ou ____________.** **Ao encontrar inconsistências no NDA padrão, substitua as cláusulas inconsistentes pelas cláusulas equivalentes do NDA cliente, exceto nas cláusulas de multas. Remova completamente qualquer cláusula de multas do NDA padrão.** Destaque todas as palavras e seções modificadas no novo documento em amarelo. Após apresentar o NDA ajustado, forneça uma lista clara de todas as inconsistências encontradas entre os dois documentos.\n\nNDA Padrão:\n{standard_nda_content}\n\nNDA Cliente:\n{user_nda_content}"}
        #{"role": "user", "content": f"Compare o NDA padrão com o NDA do cliente para gerar um novo NDA ajustado. Preencha todos os campos vazios e **não deixe campos com XXXXXXXX ou ____________.** **Ao encontrar inconsistências no NDA padrão, substitua as cláusulas inconsistentes pelas cláusulas equivalentes do NDA cliente, exceto nas cláusulas de multas. Remova completamente qualquer cláusula de multas do NDA padrão.** Destaque todas as palavras e seções modificadas no novo documento em amarelo. Após apresentar o NDA ajustado, forneça uma lista clara de todas as inconsistências encontradas entre os dois documentos.\n\nNDA Padrão:\n{standard_nda_content}\n\nNDA Cliente:\n{user_nda_content}"}
        {"role": "user", "content": f"Compare o NDA padrão com o NDA do cliente para gerar um novo NDA ajustado. Preencha todos os campos vazios e **não deixe campos com XXXXXXXX ou ____________.** **Ao encontrar inconsistências no NDA padrão, copie e cole as cláusulas originais do NDA Cliente diretamente no novo documento, exceto nas cláusulas de multas. Remova completamente qualquer cláusula de multas do NDA padrão.** Destaque todas as palavras e seções modificadas no novo documento em amarelo. Após apresentar o NDA ajustado, forneça uma lista clara de todas as inconsistências encontradas entre os dois documentos.\n\nNDA Padrão:\n{standard_nda_content}\n\nNDA Cliente:\n{user_nda_content}"}
]

    response = openai.ChatCompletion.create(
        model="gpt-4o",  # Pode ser gpt-4 ou outro modelo apropriado, gpt-4o, gpt-4o-mini, gpt-3.5-turbo
        messages=messages,
        max_tokens=4000,             #gpt-4 problema de tokens, gpt-4o-mini use 5k, gpt-4o use 4k, ,
        temperature=0.5
    )

    # Processar a resposta da OpenAI
    result_text = response['choices'][0]['message']['content'].strip()

    # Criar o novo documento ajustado em DOCX baseado no NDA padrão
    doc = docx.Document()

    # Adicionar imagem no cabeçalho, se houver uma imagem extraída
    if images:
        header = doc.sections[0].header
        header_para = header.paragraphs[0]
        run = header_para.add_run()

        # Adicionar a imagem extraída no cabeçalho
        for image in images:
            run.add_picture(image, width=Inches(1.5), height=Inches(0.75))

    # Adicionar título "NDA Ajustado Baseado no NDA Padrão"
    # doc.add_heading('NDA Ajustado Baseado no NDA Padrão', 0)
    
    # Adicionar o conteúdo do NDA ajustado
    for line in result_text.split('\n'):
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # Define o alinhamento para justificado
        if "**" in line:  # Supondo que as palavras modificadas sejam marcadas com ** para realce
            parts = line.split("**")
            for i, part in enumerate(parts):
                run = para.add_run(part)
                set_font(run)  # Aplica a fonte ao texto
                if i % 2 == 1:  # As palavras entre ** serão realçadas
                    highlight_text(para, part)
        else:
            run = para.add_run(line)
            set_font(run)  # Aplica a fonte ao texto

    # Adicionar título "NDA Ajustado com Inconformidades Realçadas"
    # doc.add_heading('NDA Ajustado com Inconformidades Realçadas', 0)
    
    # Adicionar as palavras modificadas em amarelo
    # doc.add_heading('Palavras Modificadas em Amarelo', 2)

    # Salvar o documento em memória para download
    doc_io = BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)

    return send_file(doc_io, as_attachment=True, download_name='nda_ajustado.docx')

if __name__ == '__main__':
    if not os.path.exists(UPLOAD_FOLDER):
        os.makedirs(UPLOAD_FOLDER)
    app.run(debug=True)
